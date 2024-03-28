import json
from rest_framework_simplejwt.tokens import RefreshToken
from rest_framework.response import Response
from rest_framework import status
from rest_framework.decorators import api_view
from django.contrib.auth import authenticate
import pdfplumber
from docx import Document
from django.template import loader
import os
import pdfkit
from django.http import FileResponse
from urllib.parse import quote
import secrets

import fitz  # PyMuPDF
import string
def generate_random_string(length):
    alphabet = string.ascii_letters + string.digits
    return ''.join(secrets.choice(alphabet) for _ in range(length))

#def replace_placeholder(doc, placeholder, replacement):
#     for paragraph in doc.paragraphs:
#         if placeholder in paragraph.text:
#             paragraph.text = paragraph.text.replace(placeholder, replacement)
#
# def pdf_to_docx(pdf_path, docx_path, replacements):
#     # Open the PDF file
#     with pdfplumber.open(pdf_path) as pdf:
#         # Create a new Word document
#         doc = Document()
#         # Loop through each page in the PDF
#         for page in pdf.pages:
#             # Extract text and formatting from the PDF page
#             text = page.extract_text()
#             # Add extracted text to the Word document
#             doc.add_paragraph(text)
#         # Save the Word document
#         doc.save(docx_path)
#         # Open the saved document for replacing placeholders
#         doc = Document(docx_path)
#
#         # Replace placeholders with actual content
#         for placeholder, replacement in replacements.items():
#             replace_placeholder(doc, placeholder, replacement)
#         # Save the updated Word document
#         doc.save(docx_path)
#


@api_view(['POST'])
def pdf_generator_from_account_template(request):
    try:
        data=json.loads(request.body)
        combined_data = f"{data.get('account')} {data.get('dispute_reason_in_bullet_list')}"
        # replacements = {
        # '[client_first_name]': data.get('client_first_name'),
        # '[client_middle_name]': data.get('client_middle_name'),
        # '[client_last_name]': data.get('client_last_name'),
        # '[client_address]': data.get('client_address'),
        # '[client_city]': data.get('client_city'),
        # '[client_state]': data.get('client_state'),
        # '[client_postal_code]': data.get('client_postal'),
        # '[ss_number]': data.get('ss_number'),
        # '[bdate]': data.get('bdate'),
        # '[account , dispute_reason_in_bullet_list]':combined_data
        #     }
        # to be appended at the end of generated docx or pdfs
        random_characters=generate_random_string(8)
        print(random_characters)
        pdf_path = 'pdf/ACCOUNTS DOCUMENT.pdf'
        docx_path = f'generated_docx/ACCOUNTS_DOCUMENT_{random_characters}.docx'
        output_pdf_path=f'generated_pdf/ACCOUNTS_DOCUMENT_{random_characters}.pdf'

        context = {
        'client_first_name': data.get('client_first_name'),
        'client_middle_name': data.get('client_middle_name'),
        'client_last_name': data.get('client_last_name'),
        'client_address': data.get('client_address'),
        'client_city': data.get('client_city'),
        'client_state': data.get('client_state'),
        'client_postal_code': data.get('client_postal'),
        'ss_number': data.get('ss_number'),
        'bdate': data.get('bdate'),
        'account':data.get('account'),
        'dispute_reason_in_bullet_list':data.get('dispute_reason_in_bullet_list'),
            }
        pdf_path=template_to_pdf(context,"ACCOUNTS_DOCUMENT.html",output_pdf_path,docx_path)
        print(pdf_path)
        return Response({'message':"Sucessfull"},status=200)
    except Exception as e:
        error=str(e)
        print(error)
        return Response({'error':f"Unexpected error occured {error}"},status=400)

@api_view(['POST'])
def pdf_generator_from_all_purpose_credit_template(request):
    try:
        data=json.loads(request.body)
        # replacements = {
        # '[client_first_name]': data.get('client_first_name'),
        # '[client_middle_name]': data.get('client_middle_name'),
        # '[client_last_name]': data.get('client_last_name'),
        # '[client_address]': data.get('client_address'),
        # '[client_city]': data.get('client_city'),
        # '[client_state]': data.get('client_state'),
        # '[client_postal_code]': data.get('client_postal_code'),
        # '[ss_number]': data.get('ss_number'),
        # '[bdate]': data.get('bdate'),
        # '[disputed_personal_Information]':data.get('disputed_personal_Information'),
        # '[disputed_personal_Information_Instruction]':data.get('disputed_personal_Information_Instruction'),
        # '[disputed_accounts]':data.get('disputed_accounts'),
        # '[disputed_accounts_Instruction]':data.get('disputed_accounts_Instruction'),
        # '[disputed_Inquiries]':data.get('disputed_Inquiries'),
        # '[disputed_Inquiries_Instruction]':data.get('disputed_Inquiries_Instruction'),
        #
        # }
        random_characters=generate_random_string(8)
        print(random_characters)

        pdf_path = 'ALL PURPOSE CREDIT DOCUMENT.pdf'
        docx_path = f'generated_docx/ALL_PURPOSE_CREDIT_DOCUMENT_{random_characters}.docx'
        output_pdf_path=f'generated_pdf/ALL_PURPOSE_CREDIT_DOCUMENT_{random_characters}.pdf'

        context = {
        'client_first_name': data.get('client_first_name'),
        'client_middle_name': data.get('client_middle_name'),
        'client_last_name': data.get('client_last_name'),
        'client_address': data.get('client_address'),
        'client_city': data.get('client_city'),
        'client_state': data.get('client_state'),
        'client_postal_code': data.get('client_postal'),
        'ss_number': data.get('ss_number'),
        'bdate': data.get('bdate'),
        'disputed_personal_Information':data.get('disputed_personal_Information'),
        'disputed_personal_Information_Instruction':data.get('disputed_personal_Information_Instruction'),
        'disputed_accounts':data.get('disputed_accounts'),
        'disputed_accounts_Instruction':data.get('disputed_accounts_Instruction'),
        'disputed_Inquiries':data.get('disputed_Inquiries'),
        'disputed_Inquiries_Instruction':data.get('disputed_Inquiries_Instruction')
        }
        pdf_path=template_to_pdf(context,"templates/ALL_PURPOSE_CREDIT_DOCUMENT.html",output_pdf_path,docx_path)
        print(pdf_path)
        return Response({'message':"Sucessfull"},status=200)
    except Exception as e:
        error=str(e)
        print(error)
        return Response({'error':f"Unexpected error occured {error}"},status=400)

@api_view(['POST'])
def pdf_generator_from_bankruptcy_template(request):
    try:
        data=json.loads(request.body)
        # replacements = {
        # '[client_first_name]': data.get('client_first_name'),
        # '[client_middle_name]': data.get('client_middle_name'),
        # '[client_last_name]': data.get('client_last_name'),
        # '[client_address]': data.get('client_address'),
        # '[client_city]': data.get('client_city'),
        # '[client_state]': data.get('client_state'),
        # '[client_postal_code]': data.get('client_postal'),
        # '[ss_number]': data.get('ss_number'),
        # '[bdate]': data.get('bdate'),
        # '[bankruptcy]':data.get('bankruptcy'),
        # '[bankruptcy_reason]':data.get('bankruptcy_reason'),
        # '[bankruptcy_instruction]':data.get('bankruptcy_instruction'),
        #     }
        random_characters=generate_random_string(8)
        print(random_characters)
        pdf_path = 'pdf/BANKRUPTCY DOCUMENT.pdf'
        docx_path = f'generated_docx/BANKRUPTCY_DOCUMENT_{random_characters}.docx'
        output_pdf_path=f'generated_pdf/BANKRUPTCY_DOCUMENT_{random_characters}.pdf'

        context = {
        'client_first_name': data.get('client_first_name'),
        'client_middle_name': data.get('client_middle_name'),
        'client_last_name': data.get('client_last_name'),
        'client_address': data.get('client_address'),
        'client_city': data.get('client_city'),
        'client_state': data.get('client_state'),
        'client_postal_code': data.get('client_postal'),
        'ss_number': data.get('ss_number'),
        'bdate': data.get('bdate'),
        'bankruptcy':data.get('bankruptcy'),
        'bankruptcy_reason':data.get('bankruptcy_reason'),
        'bankruptcy_instruction':data.get('bankruptcy_instruction'),
            }
        pdf_path=template_to_pdf(context,"templates/BANKRUPTCY_DOCUMENT.html",output_pdf_path,docx_path)
        print(pdf_path)
        return Response({'message':"Sucessfull"},status=200)
    except Exception as e:
        error=str(e)
        print(error)
        return Response({'error':f"Unexpected error occured {error}"},status=400)

@api_view(['POST'])
def pdf_generator_from_collection_template(request):
    try:
        data=json.loads(request.body)
        combined_data = f"{data.get('account')} {data.get('dispute_reason_in_bullet_list')}"
        # replacements = {
        # '[client_first_name]': data.get('client_first_name'),
        # '[client_middle_name]': data.get('client_middle_name'),
        # '[client_last_name]': data.get('client_last_name'),
        # '[client_address]': data.get('client_address'),
        # '[client_city]': data.get('client_city'),
        # '[client_state]': data.get('client_state'),
        # '[client_postal_code]': data.get('client_postal'),
        # '[ss_number]': data.get('ss_number'),
        # '[bdate]': data.get('bdate'),
        # '[disputed_collection]':data.get('disputed_collection'),
        # '[disputed_collection_Instruction]':data.get('disputed_collection_Instruction'),
        #     }
        random_characters=generate_random_string(8)
        print(random_characters)
        pdf_path = 'pdf/COLLECTION DOCUMENT.pdf'
        docx_path = f'generated_docx/COLLECTION_DOCUMENT_{random_characters}.docx'
        output_pdf_path=f'generated_docx/COLLECTION_DOCUMENT_{random_characters}.pdf'
        # command = "abiword --to=pdf output.docx output.pdf"
        # os.system(command)
        context = {
        'client_first_name': data.get('client_first_name'),
        'client_middle_name': data.get('client_middle_name'),
        'client_last_name': data.get('client_last_name'),
        'client_address': data.get('client_address'),
        'client_city': data.get('client_city'),
        'client_state': data.get('client_state'),
        'client_postal_code': data.get('client_postal'),
        'ss_number': data.get('ss_number'),
        'bdate': data.get('bdate'),
        'account':data.get('account'),
        'disputed_collection':data.get('disputed_collection'),
        'disputed_collection_Instruction':data.get('disputed_collection_Instruction'),
            }
        pdf_path=template_to_pdf(context,"templates/COLLECTION_DOCUMENT.html",output_pdf_path,docx_path)
        return Response({'message':"Sucessfull"},status=200)
    except Exception as e:
        error=str(e)
        print(error)
        return Response({'error':f"Unexpected error occured {error}"},status=400)

@api_view(['POST'])
def pdf_generator_from_identity_theft_affidavit_template(request):
    try:
        data=json.loads(request.body)
        # replacements = {
        # '[client_first_name]': data.get('client_first_name'),
        # '[client_middle_name]': data.get('client_middle_name'),
        # '[client_last_name]': data.get('client_last_name'),
        # '[client_address]': data.get('client_address'),
        # '[client_city]': data.get('client_city'),
        # '[client_state]': data.get('client_state'),
        # '[client_postal_code]': data.get('client_postal'),
        # '[ss_number]': data.get('ss_number'),
        # '[bdate]': data.get('bdate'),
        # '[id_number]':data.get('id_number'),
        # '[time_at_address]':data.get('time_at_address'),
        # '[todays_current_date]':data.get('todays_current_date'),
        # '[amount_spent_fixing_credit]':data.get('amount_spent_fixing_credit'),
        # '[date_identity_theft_started]':data.get('date_identity_theft_started')
        #     }
        random_characters=generate_random_string(8)
        print(random_characters)
        pdf_path = 'pdf/IDENTITY THEFT AFFIDAVIT.pdf'
        docx_path = f'generated_docx/IDENTITY_THEFT_AFFIDAVIT_{random_characters}.docx'
        output_pdf_path=f'generated_pdf/IDENTITY_THEFT_AFFIDAVIT_{random_characters}.pdf'

        context = {
        'client_first_name': data.get('client_first_name'),
        'client_middle_name': data.get('client_middle_name'),
        'client_last_name': data.get('client_last_name'),
        'client_address': data.get('client_address'),
        'client_city': data.get('client_city'),
        'client_state': data.get('client_state'),
        'client_postal_code': data.get('client_postal'),
        'ss_number': data.get('ss_number'),
        'bdate': data.get('bdate'),
        'id_number':data.get('id_number'),
        'time_at_address':data.get('time_at_address'),
        'todays_current_date':data.get('todays_current_date'),
        'amount_spent_fixing_credit':data.get('amount_spent_fixing_credit'),
        'date_identity_theft_started':data.get('date_identity_theft_started')
            }
        pdf_path=template_to_pdf(context,"templates/IDENTITY_THEFT_AFFIDAVIT.html",output_pdf_path,docx_path)
        return Response({'message':"Sucessfull"},status=200)
    except Exception as e:
        error=str(e)
        print(error)
        return Response({'error':f"Unexpected error occured {error}"},status=400)

@api_view(['POST'])
def pdf_generator_from_identity_theft_fcra_template(request):
    try:
        data=json.loads(request.body)
        # replacements = {
        # '[client_first_name]': data.get('client_first_name'),
        # '[client_middle_name]': data.get('client_middle_name'),
        # '[client_last_name]': data.get('client_last_name'),
        # '[client_address]': data.get('client_address'),
        # '[client_city]': data.get('client_city'),
        # '[client_state]': data.get('client_state'),
        # '[client_postal_code]': data.get('client_postal'),
        # '[ss_number]': data.get('ss_number'),
        # '[bdate]': data.get('bdate'),
        #     }
        random_characters=generate_random_string(8)
        print(random_characters)
        pdf_path = 'pdf/IDENTITY THEFT LETTER FCRA .pdf'
        docx_path = f'generated_docx/IDENTITY_THEFT_LETTER_FCRA_{random_characters}.docx'
        output_pdf_path=f'generated_pdf/IDENTITY_THEFT_LETTER_FCRA_{random_characters}.pdf'
        # command = "abiword --to=pdf output.docx output.pdf"
        # os.system(command)
        context = {
        'client_first_name': data.get('client_first_name'),
        'client_middle_name': data.get('client_middle_name'),
        'client_last_name': data.get('client_last_name'),
        'client_address': data.get('client_address'),
        'client_city': data.get('client_city'),
        'client_state': data.get('client_state'),
        'client_postal_code': data.get('client_postal'),
        'ss_number': data.get('ss_number'),
        'bdate': data.get('bdate'),
            }
        pdf_path=template_to_pdf(context,"templates/IDENTITY_THEFT_LETTER_FCRA.html",output_pdf_path,docx_path)
        return Response({'message':"Sucessfull"},status=200)
    except Exception as e:
        error=str(e)
        print(error)
        return Response({'error':f"Unexpected error occured {error}"},status=400)

@api_view(['POST'])
def pdf_generator_from_inquiry_template(request):
    try:
        data=json.loads(request.body)
        # replacements = {
        # '[client_first_name]': data.get('client_first_name'),
        # '[client_middle_name]': data.get('client_middle_name'),
        # '[client_last_name]': data.get('client_last_name'),
        # '[client_address]': data.get('client_address'),
        # '[client_city]': data.get('client_city'),
        # '[client_state]': data.get('client_state'),
        # '[client_postal_code]': data.get('client_postal'),
        # '[ss_number]': data.get('ss_number'),
        # '[bdate]': data.get('bdate'),
        # '[disputed_Inquiries]':data.get('disputed_Inquiries'),
        # '[disputed_Inquiries_date]':data.get('disputed_Inquiries_date'),
        # '[disputed_Inquiries_reason]':data.get('disputed_Inquiries_reason'),
        # '[disputed_Inquiries_Instruction]':data.get('disputed_Inquiries_Instruction')
        #     }
        random_characters=generate_random_string(8)
        print(random_characters)
        pdf_path = 'pdf/INQUIRY DOCUMENT.pdf'
        docx_path = f'generated_docx/INQUIRY_DOCUMENT_{random_characters}.docx'
        output_pdf_path=f'generated_pdf/INQUIRY_DOCUMENT_{random_characters}.pdf'

        context = {
        'client_first_name': data.get('client_first_name'),
        'client_middle_name': data.get('client_middle_name'),
        'client_last_name': data.get('client_last_name'),
        'client_address': data.get('client_address'),
        'client_city': data.get('client_city'),
        'client_state': data.get('client_state'),
        'client_postal_code': data.get('client_postal'),
        'ss_number': data.get('ss_number'),
        'bdate': data.get('bdate'),
        'disputed_Inquiries':data.get('disputed_Inquiries'),
        'disputed_Inquiries_date':data.get('disputed_Inquiries_date'),
        'disputed_Inquiries_reason':data.get('disputed_Inquiries_reason'),
        'disputed_Inquiries_Instruction':data.get('disputed_Inquiries_Instruction')            }
        pdf_path=template_to_pdf(context,"templates/INQUIRY_DOCUMENT.html",output_pdf_path,docx_path)
        print(pdf_path)
        return Response({'message':"Sucessfull"},status=200)
    except Exception as e:
        error=str(e)
        print(error)
        return Response({'error':f"Unexpected error occured {error}"},status=400)

@api_view(['POST'])
def pdf_generator_from_personal_information_template(request):
    try:
        data=json.loads(request.body)
        # replacements = {
        # '[client_first_name]': data.get('client_first_name'),
        # '[client_middle_name]': data.get('client_middle_name'),
        # '[client_last_name]': data.get('client_last_name'),
        # '[client_address]': data.get('client_address'),
        # '[client_city]': data.get('client_city'),
        # '[client_state]': data.get('client_state'),
        # '[client_postal_code]': data.get('client_postal'),
        # '[ss_number]': data.get('ss_number'),
        # '[bdate]': data.get('bdate'),
        # '[client_sign]':data.get('client_sign'),
        # '[personal_Information_explanation]':data.get('personal_Information_explanation'),
        # '[employer_Information_explanation]':data.get('employer_Information_explanation'),
        #     }
        random_characters=generate_random_string(8)
        print(random_characters)
        pdf_path = 'pdf/PERSONAL INFORMATION DOCUMENT.pdf'
        docx_path = f'generated_docx/PERSONAL_INFORMATION_DOCUMENT_{random_characters}.docx'
        output_pdf_path=f'generated_pdf/PERSONAL_INFORMATION_DOCUMENT_{random_characters}.pdf'

        context = {
        'client_first_name': data.get('client_first_name'),
        'client_middle_name': data.get('client_middle_name'),
        'client_last_name': data.get('client_last_name'),
        'client_address': data.get('client_address'),
        'client_city': data.get('client_city'),
        'client_state': data.get('client_state'),
        'client_postal_code': data.get('client_postal'),
        'ss_number': data.get('ss_number'),
        'bdate': data.get('bdate'),
         'client_sign':data.get('client_sign'),
        'personal_Information_explanation':data.get('personal_Information_explanation'),
        'employer_Information_explanation':data.get('employer_Information_explanation'),          }
        pdf_path=template_to_pdf(context,"templates/PERSONAL_INFORMATION_DOCUMENT.html",output_pdf_path,docx_path)
        print(pdf_path)
        return Response({'message':"Sucessfull"},status=200)
    except Exception as e:
        error=str(e)
        print(error)
        return Response({'error':f"Unexpected error occured {error}"},status=400)

@api_view(['GET'])
def download_file(request,filename):
    try:
            # Return a response to download the pdf file
        response = FileResponse(open(filename, 'rb'))
        response['Content-Disposition'] = 'attachment; filename="%s"' % quote(filename)

        return response
    except Exception as e:
        error=str(e)
        print(error)
        return Response({'error':f"unexpected errror{error}"},status=400)

def pdf_to_docx(pdf_path, docx_path):
    # Open the PDF file
    with pdfplumber.open(pdf_path) as pdf:
        # Create a new Word document
        doc = Document()
        # Loop through each page in the PDF
        for page in pdf.pages:
            # Extract text and formatting from the PDF page
            text = page.extract_text()
            # Add extracted text to the Word document
            doc.add_paragraph(text)
        # Save the Word document
        doc.save(docx_path)
        # Open the saved document for replacing placeholders
        doc = Document(docx_path)

        # Replace placeholders with actual content
        # Save the updated Word document
        doc.save(docx_path)



#For Creating template to pdf based on context
def template_to_pdf(context,template,output_pdf_path,output_docx_path):
    template = loader.get_template(template)
    html_content = template.render(context)

    # Generate PDF from the HTML content
    pdf_file = pdfkit.from_string(html_content, False)
    file_path=output_pdf_path
    with open(file_path, 'wb') as f:
        f.write(pdf_file)

    pdf_to_docx(file_path,output_docx_path)
    # Create a response with the PDF content

    return True
