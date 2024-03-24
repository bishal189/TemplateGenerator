import pdfplumber
from docx import Document

def replace_placeholder(doc, placeholder, replacement):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, replacement)

def pdf_to_docx(pdf_path, docx_path, replacements):
    # Open the PDF file
    with pdfplumber.open(pdf_path) as pdf:
        # Create a new Word document
        doc = Document()

        # Loop through each page in the PDF
        for page in pdf.pages:
            # Extract text and formatting from the PDF page
            text = page.extract_text()

            # Add extracted text to the Word document
            doc.add_paragraph(text)

        # Save the Word document
        doc.save(docx_path)

        # Open the saved document for replacing placeholders
        doc = Document(docx_path)

        # Replace placeholders with actual content
        for placeholder, replacement in replacements.items():
            replace_placeholder(doc, placeholder, replacement)

        # Save the updated Word document
        doc.save(docx_path)

# Example usage
pdf_path = 'template.pdf'
docx_path = 'output.docx'
replacements = {'[client_first_name]': 'John', '[client_last_name]': 'Doe'}  # Add all your replacements here
pdf_to_docx(pdf_path, docx_path, replacements)
